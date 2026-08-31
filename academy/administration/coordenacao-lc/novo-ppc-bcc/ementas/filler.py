"""Fill MODELO EMENTA SIGAA.docx into one per-discipline doc.
Usage: filler.py <content.json> <out.docx>
Preserves the SIGAA table template (nested in sdt); fills cells by label match.
"""
import sys, json, copy
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

MODELO = str(__import__("pathlib").Path(__file__).resolve().parents[5]
              / "Downloads/workspace-drive/MODELO EMENTA SIGAA.docx")

# inline: value appended as a run to the label paragraph (same line)
INLINE = {
    "COMPONENTE CURRICULAR:": "componente",
    "CÓDIGO:": "codigo",
    "PERÍODO A SER OFERTADO:": "periodo",
    "PRÉ-REQUISITO:": "pre_requisito",
    "CORREQUISITO:": "correquisito",
    "EQUIVALÊNCIA:": "equivalencia",
    "CARGA HORÁRIA TOTAL:": "carga_total",
}
# block: value added as new paragraph(s) in the cell, after the label
BLOCK = {
    "EMENTA:": "ementa",
    "OBJETIVOS:": "objetivos",
    "CONTEÚDO PROGRAMÁTICO:": "conteudo",
    "BIBLIOGRAFIA BÁSICA:": "bibliografia_basica",
    "BIBLIOGRAFIA COMPLEMENTAR:": "bibliografia_complementar",
}


def cell_label(tc):
    txts = [t.text or "" for t in tc.findall(".//" + qn("w:t"))]
    return "".join(txts).strip()


def fill(content, out):
    d = Document(MODELO)
    tbl_el = d.element.body.findall(".//" + qn("w:tbl"))[0]
    new_el = copy.deepcopy(tbl_el)

    nd = Document()
    # drop default empty paragraph
    for p in list(nd.element.body):
        if p.tag == qn("w:p"):
            nd.element.body.remove(p)
    sectPr = nd.element.body.find(qn("w:sectPr"))
    nd.element.body.insert(list(nd.element.body).index(sectPr), new_el)
    tbl = Table(new_el, nd)

    for row in tbl.rows:
        seen = set()
        for cell in row.cells:
            if id(cell._tc) in seen:   # merged cells repeat per grid column
                continue
            seen.add(id(cell._tc))
            lab = cell_label(cell._tc)
            for key, field in INLINE.items():
                if lab.startswith(key) and content.get(field):
                    # carga default already 60h in template — only override if given
                    cell.paragraphs[0].add_run(" " + str(content[field]))
                    break
            for key, field in BLOCK.items():
                if lab.startswith(key) and content.get(field):
                    for line in str(content[field]).split("\n"):
                        cell.add_paragraph(line)
                    break
    nd.save(out)
    print("saved:", out)


if __name__ == "__main__":
    content = json.load(open(sys.argv[1], encoding="utf-8"))
    fill(content, sys.argv[2])
