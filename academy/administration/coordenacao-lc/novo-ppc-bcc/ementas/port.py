"""Port one discipline into a MODELO-SIGAA doc.

VERBATIM: every existing field value is copied byte-for-byte from the source
.docx cell (no retyping, no reformatting). Only OBJETIVOS may be supplied
separately (the sole generated field) via --objetivos-file.

Usage: port.py <source.docx> "<DISCIPLINE NAME>" <out.docx> [objetivos.txt]
"""
import sys
import copy
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

MODELO = str(__import__("pathlib").Path(__file__).resolve().parents[5]
              / "Downloads/workspace-drive/MODELO EMENTA SIGAA.docx")

# label prefix -> field key. Order matters (longest/specific first not needed; exact prefixes).
LABELS = [
    ("COMPONENTE CURRICULAR:", "componente"),
    ("CÓDIGO:", "codigo"),
    ("PERÍODO A SER OFERTADO:", "periodo"),
    ("CARGA HORÁRIA TOTAL:", "carga"),
    ("PRÉ-REQUISITO:", "pre_requisito"),
    ("CORREQUISITO:", "correquisito"),
    ("EQUIVALÊNCIA:", "equivalencia"),
    ("EMENTA:", "ementa"),
    ("OBJETIVOS:", "objetivos"),
    ("CONTEÚDO PROGRAMÁTICO:", "conteudo"),
    ("BIBLIOGRAFIA BÁSICA:", "bib_basica"),
    ("BIBLIOGRÁFIA BÁSICA:", "bib_basica"),
    ("BIBLIOGRAFIA COMPLEMENTAR:", "bib_comp"),
    ("BIBLIOGRÁFIA COMPLEMENTAR:", "bib_comp"),
    ("Teórica:", "carga_breakdown_raw"),  # whole "Teórica:...Prática:...Extensão:...EaD:" blob (one cell)
]
INLINE = {"componente", "codigo", "periodo", "carga", "pre_requisito",
          "correquisito", "equivalencia"}
BLOCK = {"ementa", "objetivos", "conteudo", "bib_basica", "bib_comp"}


def uniq_cells(row):
    seen, out = set(), []
    for c in row.cells:
        if id(c._tc) in seen:
            continue
        seen.add(id(c._tc))
        out.append(c)
    return out


def match_label(text):
    for prefix, key in LABELS:
        if text.startswith(prefix):
            return key, text[len(prefix):]   # verbatim remainder (leading space kept, lstrip later)
    return None, None


def extract(source, name):
    d = Document(source)
    for t in d.tables:
        if name.upper() in t.rows[0].cells[0].text.upper():
            fields, current = {}, None
            for row in t.rows:
                for cell in uniq_cells(row):
                    text = cell.text
                    key, rest = match_label(text)
                    if key:
                        if key == "componente" and "componente" in fields:
                            # this w:tbl packs multiple courses back-to-back (no clean
                            # 1-table-per-course boundary in the consolidated source) --
                            # stop before overwriting with the next course's rows.
                            return fields
                        current = key
                        fields[key] = rest.lstrip("\n ")
                    elif current in BLOCK:
                        # continuation cell (e.g. CONTEÚDO spanning multiple columns/rows)
                        fields[current] = (fields.get(current, "") + "\n" + text).strip("\n")
            return fields
    raise SystemExit(f"discipline not found in source: {name}")


def set_font(run, name="Times New Roman"):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(a), name)


def fill(fields, out):
    d = Document(MODELO)
    tbl_el = copy.deepcopy(d.element.body.findall(".//" + qn("w:tbl"))[0])
    nd = Document()
    for p in list(nd.element.body):
        if p.tag == qn("w:p"):
            nd.element.body.remove(p)
    sect = nd.element.body.find(qn("w:sectPr"))
    nd.element.body.insert(list(nd.element.body).index(sect), tbl_el)
    tbl = Table(tbl_el, nd)

    for row in tbl.rows:
        for cell in uniq_cells(row):
            lab = cell.text.strip()
            for prefix, key in LABELS:
                if lab.startswith(prefix) and fields.get(key):
                    val = fields[key]
                    if key == "carga_breakdown_raw":
                        # whole "Teórica:.../Prática:.../Extensão:.../EaD:..." blob lives in
                        # one cell as multiple paragraphs -- template ships it pre-filled
                        # ("60h" default), so replace every paragraph, don't append.
                        lines = (prefix + val).split("\n")
                        for extra_p in list(cell.paragraphs[1:]):
                            extra_p._element.getparent().remove(extra_p._element)
                        first_para = cell.paragraphs[0]
                        for run_ in list(first_para.runs):
                            run_._element.getparent().remove(run_._element)
                        r = first_para.add_run(lines[0])
                        set_font(r)
                        for line in lines[1:]:
                            p = cell.add_paragraph()
                            r = p.add_run(line)
                            set_font(r)
                    elif key in INLINE:
                        # replace, don't append -- some template cells ship with a
                        # baked-in default (e.g. "60h", "NÃO TEM") after the label
                        para = cell.paragraphs[0]
                        for run_ in list(para.runs):
                            run_._element.getparent().remove(run_._element)
                        r = para.add_run(prefix + " " + val)
                        set_font(r)
                    else:
                        for line in val.split("\n"):
                            p = cell.add_paragraph()
                            r = p.add_run(line)
                            set_font(r)
                    break
    # normalize EVERY run to Times New Roman (kills inherited Cambria in Docs)
    for run_el in tbl_el.iter(qn("w:r")):
        rPr = run_el.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(a), "Times New Roman")
    nd.save(out)


if __name__ == "__main__":
    src, name, out = sys.argv[1], sys.argv[2], sys.argv[3]
    fields = extract(src, name)
    if len(sys.argv) > 4:  # objetivos override (generated field)
        fields["objetivos"] = open(sys.argv[4], encoding="utf-8").read().strip()
    # report which fields were empty in source (candidates for manual fill)
    empties = [k for k in ("ementa", "objetivos", "conteudo", "bib_basica", "bib_comp")
               if not fields.get(k, "").strip()]
    fill(fields, out)
    print("saved:", out)
    print("EMPTY-IN-SOURCE:", empties)
